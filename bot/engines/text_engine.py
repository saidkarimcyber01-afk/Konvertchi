from __future__ import annotations

from io import BytesIO
from typing import Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from bot.models.documents import TextAlignment, TextStyle


ALIGNMENT_MAP = {
    TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PDF_ALIGNMENT_MAP = {
    TextAlignment.LEFT: TA_LEFT,
    TextAlignment.CENTER: TA_CENTER,
    TextAlignment.RIGHT: TA_RIGHT,
    TextAlignment.JUSTIFY: TA_JUSTIFY,
}


def _apply_style(run, style: TextStyle) -> None:
    run.bold = style in {TextStyle.BOLD, TextStyle.BOLD_ITALIC}
    run.italic = style in {TextStyle.ITALIC, TextStyle.BOLD_ITALIC}


def text_to_docx(title: str, body: str, alignment: TextAlignment, style: TextStyle, font_size: int) -> bytes:
    document = Document()
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(font_size + 2)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    paragraph = document.add_paragraph()
    paragraph.alignment = ALIGNMENT_MAP[alignment]
    run = paragraph.add_run(body)
    _apply_style(run, style)
    run.font.size = Pt(font_size)

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def text_to_pdf(title: str, body: str, alignment: TextAlignment, style: TextStyle, font_size: int) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    alignment_value = PDF_ALIGNMENT_MAP[alignment]

    base_style = ParagraphStyle(
        name="Body",
        fontSize=font_size,
        leading=font_size * 1.4,
        alignment=alignment_value,
    )
    if style in {TextStyle.BOLD, TextStyle.BOLD_ITALIC}:
        base_style.fontName = "Helvetica-Bold"
    if style in {TextStyle.ITALIC, TextStyle.BOLD_ITALIC}:
        base_style.fontName = "Helvetica-Oblique" if style == TextStyle.ITALIC else "Helvetica-BoldOblique"

    elements = [
        Paragraph(f"<b>{title}</b>", ParagraphStyle(name="Title", fontSize=font_size + 2, alignment=TA_CENTER)),
        Spacer(1, font_size),
        Paragraph(body.replace("\n", "<br />"), base_style),
    ]
    doc.build(elements)
    return buffer.getvalue()


def build_text_document(
    title: str,
    body: str,
    alignment: TextAlignment,
    style: TextStyle,
    font_size: int,
    output_format: str,
) -> Tuple[bytes, str]:
    if output_format == "docx":
        return text_to_docx(title, body, alignment, style, font_size), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if output_format == "pdf":
        return text_to_pdf(title, body, alignment, style, font_size), "application/pdf"
    raise ValueError("Unsupported output format")
